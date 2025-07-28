import streamlit as st
import pandas as pd
import io
from docx import Document
from datetime import datetime
from fpdf import FPDF
import re
import traceback

st.set_page_config(page_title="Payroll Summary Generator", layout="wide")

# ——— Simple login gating ———
VALID_USERS = {
    "alice": "wonderland123",
    "bob":   "builder456",
    # add your users here
}

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if not st.session_state.logged_in:
    st.title("🔒 Please log in")
    with st.form("login_form"):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Log in")
        if submitted:
            if VALID_USERS.get(username) == password:
                st.session_state.logged_in = True
                st.experimental_rerun()  # rerun to enter the app
            else:
                st.error("❌ Incorrect username or password")
    st.stop()


# ——— helper to sum normalized column names ———
def sum_norm(df, *target_names):
    """
    For each name in target_names, normalize (strip non-letters, lowercase)
    and look for a matching column in df. Return the .sum() of the first match.
    """
    for target in target_names:
        tgt_key = re.sub(r"[^a-z]", "", target.lower())
        for col in df.columns:
            col_key = re.sub(r"[^a-z]", "", col.lower())
            if col_key == tgt_key:
                return df[col].sum()
    return 0

st.title("📊 Payroll Summary Generator")

uploaded_file = st.file_uploader(
    "📂 Upload raw payroll report (Excel)",
    type=["xlsx"],
    help="Select any .xlsx file with the expected columns"
)

if not uploaded_file:
    st.info("Please upload an Excel file above to proceed.")
    st.stop()

# ——— load & detect header ———
xlsx    = pd.ExcelFile(uploaded_file)
sheet   = st.selectbox("Select a sheet to process", xlsx.sheet_names)
preview = pd.read_excel(uploaded_file, sheet_name=sheet, header=None)
hdr_row = preview[preview.eq("No.").any(axis=1)].index[0]
df_raw  = pd.read_excel(uploaded_file, sheet_name=sheet, header=hdr_row)

# ——— clean column names ———
df_raw.columns = (
    df_raw.columns
          .str.replace("`", "", regex=False)
          .str.replace("’", "", regex=False)
          .str.replace("‘", "", regex=False)
          .str.strip()
)
df_raw = (
    df_raw
    .rename(columns={"EPFEE": "EPF", "SocEE": "Socso", "EISEE": "EIS"})
    .loc[:, ~df_raw.columns.duplicated()]
    .reset_index(drop=True)
)

st.subheader("Preview of selected sheet")
st.dataframe(df_raw.head())
st.write("🧲 Columns in file:", list(df_raw.columns))

# ——— find cost‑center column ———
def match_cost_center_column(columns):
    acceptable = ["C/Center","Cost Center","Center","C Center","C-Center"]
    norm_map    = {col: re.sub(r"[^a-z]", "", col.lower()) for col in columns}
    for tgt in acceptable:
        t_norm = re.sub(r"[^a-z]", "", tgt.lower())
        for orig, norm in norm_map.items():
            if norm == t_norm:
                return orig
    return None

cost_col = match_cost_center_column(df_raw.columns)
if not cost_col:
    st.error("❌ Cost Center column not found.")
    st.stop()

depts    = df_raw[cost_col].dropna().unique().tolist()
sel_dept = st.selectbox("Select Department", depts)
df       = df_raw[df_raw[cost_col] == sel_dept].copy()

# ——— drop resigned staff ———
df["Joined"] = pd.to_datetime(df["Joined"], errors="coerce")
df["Resign"] = pd.to_datetime(df["Resign"], errors="coerce")
df = df[df["Resign"].isna()]

# ——— numeric conversion ———
ignore   = ["Name","Emp No",cost_col,"Joined","Resign"]
num_cols = df.columns.drop(ignore, errors="ignore")
df[num_cols] = df[num_cols].apply(pd.to_numeric, errors="coerce").fillna(0)

# ——— payroll summary ———
df["Gross Pay"] = (
      df.get("M/Basic", 0)
    + df.get("OT Amt 1½", 0)
    + df.get("MEC", 0)
    + df.get("ALL", 0)
    + df.get("OVT", 0)
    + df.get("MS", 0) + df.get("NS", 0) + df.get("ICP", 0)
    + df.get("BAC", 0) + df.get("BSC", 0) + df.get("BBB", 0)
    + df.get("BAL", 0) + df.get("BOT", 0) + df.get("CSN", 0)
)
df["EPF"]             = df.get("EPF", 0)
df["Socso"]           = df.get("Socso", 0)
df["EIS"]             = df.get("EIS", 0)
df["PCB"]             = df.get("PCB", 0)
df["Total Deduction"] = df[["EPF","Socso","EIS","PCB"]].sum(axis=1)
df["Net Pay"]         = df["Gross Pay"] - df["Total Deduction"]

st.subheader(f"Payroll Summary: {sel_dept}")
st.dataframe(df)

# ——— Excel summary download ———
excel_buf = io.BytesIO()
with pd.ExcelWriter(excel_buf, engine="xlsxwriter") as writer:
    df.to_excel(writer, index=False, sheet_name=sel_dept)
st.download_button(
    "📅 Download Excel Summary",
    data=excel_buf.getvalue(),
    file_name=f"payroll_summary_{sel_dept}.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
)

# ——— Invoice Section ———
st.markdown("### 📄 Generate Department Invoice")
if st.button("Generate Invoice"):
    try:
        # headcount & sums
        n               = len(df)
        gross_sum       = df["Gross Pay"].sum()
        overtime_total  = df["OT Amt 1½"].sum() + df["BOT"].sum()

        # employer statutory: look for EPF ER, EIS ER, and Socso ER variants
        epf_er   = sum_norm(df, "EPF ER",   "EPF'ER",  "EPFER")
        eis_er   = sum_norm(df, "EIS ER",   "EIS'ER",  "EISER")
        socso_er = sum_norm(
            df,
            "Socso ER",
            "SOC ER",
            "SOC'ER",
            "SOCSOER",
            "Soc 'EE"    # in case your sheet literally says "Soc 'EE"
        )
        emp_stat = epf_er + socso_er + eis_er

        hrdf_amt      = df.get("HRDF", 0).sum()
        insurance_fee = 50

        # Line 1: Wages = Gross Pay – Overtime
        wages = gross_sum - overtime_total

        # Management fee = 15% of (Wages + OT + Statutory + HRDF)
        mgmt_base = wages + overtime_total + emp_stat + hrdf_amt
        mgmt_fee  = 0.15 * mgmt_base

        date_str = datetime.today().strftime("%Y-%m-%d")

        # build invoice lines
        items = [
            {"No.":1, "Description":"Wages",                              "Qty":1, "U.Price":f"{wages:,.2f}",      "Amount":f"{wages:,.2f}"},
            {"No.":2, "Description":"Overtime",                           "Qty":1, "U.Price":f"{overtime_total:,.2f}", "Amount":f"{overtime_total:,.2f}"},
            {"No.":3, "Description":"Employer Statutory (EPF+Socso+EIS)", "Qty":1, "U.Price":f"{emp_stat:,.2f}",    "Amount":f"{emp_stat:,.2f}"},
            {"No.":4, "Description":"HRDF",                               "Qty":1, "U.Price":f"{hrdf_amt:,.2f}",    "Amount":f"{hrdf_amt:,.2f}"},
            {"No.":5, "Description":"Medical Fee (Excl. Mgmt Fee)",       "Qty":1, "U.Price":"",                     "Amount":""},
            {"No.":6, "Description":"Insurance Claim (Excl. Mgmt Fee)",   "Qty":n, "U.Price":f"{insurance_fee:,.2f}", "Amount":f"{insurance_fee*n:,.2f}"},
            {"No.":7, "Description":"15% Management Fee",                "Qty":1, "U.Price":f"{mgmt_fee:,.2f}",    "Amount":f"{mgmt_fee:,.2f}"},
        ]
        inv_df = pd.DataFrame(items)

        # preview
        st.subheader("Invoice Preview")
        st.table(inv_df)

        # SST calculations
        total_excl_sst = sum(float(x.replace(",", "")) for x in inv_df["Amount"] if x)
        sst_rate       = 0.08
        sst_amount     = total_excl_sst * sst_rate
        total_incl_sst = total_excl_sst + sst_amount

        st.write(f"**Total (Excl. SST):** RM {total_excl_sst:,.2f}")
        st.write(f"**SST @8%:** RM {sst_amount:,.2f}")
        st.write(f"**Total (Incl. SST):** RM {total_incl_sst:,.2f}")

        # — Download Word Invoice ———
        doc = Document()
        doc.add_heading("INVOICE", level=0)
        doc.add_paragraph(f"Date: {date_str}")
        doc.add_paragraph(f"Department: {sel_dept.upper()}")

        tbl = doc.add_table(rows=1, cols=len(inv_df.columns))
        for i, col in enumerate(inv_df.columns):
            tbl.rows[0].cells[i].text = col
        for _, row in inv_df.iterrows():
            cells = tbl.add_row().cells
            for i, col in enumerate(inv_df.columns):
                cells[i].text = str(row[col])

        doc.add_paragraph(f"\nTotal (Excl. SST): RM {total_excl_sst:,.2f}")
        doc.add_paragraph(f"SST @8%: RM {sst_amount:,.2f}")
        doc.add_paragraph(f"Total (Incl. SST): RM {total_incl_sst:,.2f}")

        word_buf = io.BytesIO()
        doc.save(word_buf)
        word_buf.seek(0)
        st.download_button(
            "📥 Download Word Invoice",
            data=word_buf.getvalue(),
            file_name=f"invoice_{sel_dept}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # — Download PDF Invoice ———
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial","B",16)
        pdf.cell(0,10,"INVOICE",ln=True,align="C")
        pdf.set_font("Arial","",12)
        pdf.cell(0,8,f"Date: {date_str}",ln=True)
        pdf.cell(0,8,f"Department: {sel_dept}",ln=True)
        pdf.ln(5)

        pdf.set_font("Arial","B",12)
        widths = [10,80,20,40,40]
        for w, col in zip(widths, inv_df.columns):
            pdf.cell(w,8,col,border=1)
        pdf.ln()

        pdf.set_font("Arial","",12)
        for _, row in inv_df.iterrows():
            pdf.cell(widths[0],8,str(row["No."]),border=1)
            pdf.cell(widths[1],8,row["Description"],border=1)
            pdf.cell(widths[2],8,str(row["Qty"]),border=1)
            pdf.cell(widths[3], 8, row["U.Price"], border=1)
            pdf.cell(widths[4], 8, row["Amount"],  border=1, ln=True)

            

        pdf.ln(5)
        pdf.cell(0,8,f"Total (Excl. SST): RM {total_excl_sst:,.2f}",ln=True)
        pdf.cell(0,8,f"SST @8%: RM {sst_amount:,.2f}",ln=True)
        pdf.cell(0,8,f"Total (Incl. SST): RM {total_incl_sst:,.2f}",ln=True)

        pdf_buf = io.BytesIO()
        pdf_buf.write(pdf.output(dest="S").encode("latin-1"))
        pdf_buf.seek(0)
        st.download_button(
            "📥 Download PDF Invoice",
            data=pdf_buf.getvalue(),
            file_name=f"invoice_{sel_dept}.pdf",
            mime="application/pdf"
        )

    except Exception as e:
        traceback.print_exc()
        st.error(f"Something went wrong: {e}")